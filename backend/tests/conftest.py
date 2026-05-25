import io
import os
import tempfile
from pathlib import Path

import mammoth
import pytest
from fastapi.testclient import TestClient

# Point at a temp docs dir before importing the app.
_tmp_docs = tempfile.mkdtemp()
os.environ.setdefault("DOCS_PATH", _tmp_docs)
os.environ.setdefault("OPENROUTER_API_KEY", "test-key")

from app.main import app  # noqa: E402

DOCS_DIR = Path(_tmp_docs)


@pytest.fixture(scope="session")
def docs_dir() -> Path:
    return DOCS_DIR


@pytest.fixture(scope="session")
def sample_docx(docs_dir: Path) -> Path:
    """Create a minimal valid .docx fixture using python-docx if available,
    otherwise write a tiny hand-crafted .docx zip."""
    path = docs_dir / "Sample Document.docx"
    try:
        import docx  # python-docx

        doc = docx.Document()
        doc.add_heading("Sample Heading", level=1)
        doc.add_paragraph("This is a sample paragraph with bold text.")
        doc.save(str(path))
    except ImportError:
        # Fallback: copy a known-good minimal docx from mammoth's test assets
        import zipfile

        with zipfile.ZipFile(path, "w") as zf:
            zf.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>",
            )
            zf.writestr(
                "_rels/.rels",
                '<?xml version="1.0"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                "</Relationships>",
            )
            zf.writestr(
                "word/_rels/document.xml.rels",
                '<?xml version="1.0"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
            )
            zf.writestr(
                "word/document.xml",
                '<?xml version="1.0"?>'
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:body>"
                "<w:p><w:pPr><w:pStyle w:val=\"Heading1\"/></w:pPr><w:r><w:t>Sample Heading</w:t></w:r></w:p>"
                "<w:p><w:r><w:t>This is a sample paragraph.</w:t></w:r></w:p>"
                "</w:body></w:document>",
            )
    return path


@pytest.fixture(scope="session")
def client() -> TestClient:
    return TestClient(app)
